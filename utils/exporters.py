# utils/exporters.py
#
# Export Utilities
# ----------------
# Functions to export the final proposal to downloadable file formats.
# Currently supports:
#   - .docx (Word document) via python-docx
#   - .txt  (plain text fallback — no extra dependencies)
#
# Why python-docx and not reportlab/WeasyPrint for PDF?
# For a capstone demo, Word format is more practical — evaluators and
# supervisors can open it without special software, and it's editable.
# PDF export can be added later if needed.
#
# Install: pip install python-docx
#
# Author: Krit Prakash (M.Tech Capstone — Vignan Institute)

import os
import json
import logging
from datetime import datetime

logging.basicConfig(level=logging.INFO, format="%(levelname)s | %(message)s")
logger = logging.getLogger(__name__)

# output directory for all exports
EXPORT_DIR = "data/exports"


def ensure_export_dir():
    """Create the exports directory if it doesn't exist."""
    os.makedirs(EXPORT_DIR, exist_ok=True)


def export_proposal_docx(
    proposal_draft: dict,
    budget_timeline: dict,
    eval_scores: dict,
    topic: str,
    iteration: int = 0
) -> str:
    """
    Export the final proposal as a .docx Word document.
    Includes all 6 proposal sections, a budget table, and the evaluation summary.

    Args:
        proposal_draft:  Dict with the 6 proposal sections
        budget_timeline: Dict from BudgetTimelineAgent
        eval_scores:     Dict from EvaluationAgent
        topic:           Research topic string (used in filename)
        iteration:       Final iteration number (used in filename)

    Returns:
        Path to the saved .docx file
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.warning("python-docx not installed. Falling back to .txt export.")
        return export_proposal_txt(proposal_draft, budget_timeline, eval_scores, topic)

    ensure_export_dir()

    doc = Document()

    # ── document title page ──────────────────────────────────────────────────
    title_para = doc.add_heading("Research Grant Proposal", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Topic: {topic}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Generated: {datetime.now().strftime('%B %d, %Y')}  |  "
        f"Score: {eval_scores.get('total_score', 0)}/100  |  "
        f"Iterations: {iteration}"
    ).font.size = Pt(10)

    doc.add_page_break()

    # ── table of contents (manual — no built-in TOC in python-docx) ──────────
    doc.add_heading("Contents", level=1)
    sections_order = [
        "Title", "Abstract", "Problem Statement",
        "Objectives", "Methodology", "Expected Impact"
    ]
    for i, section in enumerate(sections_order, 1):
        doc.add_paragraph(f"{i}.  {section}", style="List Number")
    doc.add_paragraph("7.  Budget & Cost Breakdown")
    doc.add_paragraph("8.  Project Timeline")
    doc.add_paragraph("9.  Evaluation Summary")
    doc.add_page_break()

    # ── proposal sections ─────────────────────────────────────────────────────
    for section_name in sections_order:
        content = proposal_draft.get(section_name, "")
        if not content:
            continue

        doc.add_heading(section_name, level=1)

        # Objectives are usually bullet-pointed — detect and format them
        if section_name == "Objectives" and "\n" in content:
            for line in content.split("\n"):
                line = line.strip()
                if line:
                    # strip common bullet characters if LLM added them
                    clean = line.lstrip("•-*123456789. ")
                    doc.add_paragraph(clean, style="List Bullet")
        else:
            doc.add_paragraph(content)

        doc.add_paragraph("")  # spacer

    doc.add_page_break()

    # ── budget table ──────────────────────────────────────────────────────────
    doc.add_heading("Budget & Cost Breakdown", level=1)

    budget_items = budget_timeline.get("budget_items", [])
    total_cost   = budget_timeline.get("total_cost_usd", 0)
    within_limit = budget_timeline.get("within_limit", True)

    if budget_items:
        # create table with header row
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        # header
        hdr_cells = table.rows[0].cells
        headers = ["Category", "Item", "Cost (USD)", "Justification"]
        for i, hdr in enumerate(headers):
            hdr_cells[i].text = hdr
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True

        # data rows
        for item in budget_items:
            row_cells = table.add_row().cells
            row_cells[0].text = item.get("category", "")
            row_cells[1].text = item.get("item", "")
            row_cells[2].text = f"${item.get('cost_usd', 0):,}"
            row_cells[3].text = item.get("justification", "")

        doc.add_paragraph("")
        total_para = doc.add_paragraph()
        total_run  = total_para.add_run(f"Total Budget: ${total_cost:,} USD")
        total_run.font.bold = True
        total_run.font.size = Pt(11)

        if not within_limit:
            warn = doc.add_paragraph()
            warn.add_run(
                "⚠ Note: Budget exceeds the stated guideline limit."
            ).font.color.rgb = RGBColor(0xFF, 0x66, 0x00)
    else:
        doc.add_paragraph("Budget data not available.")

    doc.add_page_break()

    # ── timeline ──────────────────────────────────────────────────────────────
    doc.add_heading("Project Timeline", level=1)

    milestones = budget_timeline.get("timeline_milestones", [])
    if milestones:
        tl_table = doc.add_table(rows=1, cols=3)
        tl_table.style = "Table Grid"

        th = tl_table.rows[0].cells
        for i, label in enumerate(["Month", "Milestone", "Deliverable"]):
            th[i].text = label
            th[i].paragraphs[0].runs[0].font.bold = True

        for m in milestones:
            row = tl_table.add_row().cells
            row[0].text = str(m.get("month", ""))
            row[1].text = m.get("milestone", "")
            row[2].text = m.get("deliverable", "")
    else:
        doc.add_paragraph("Timeline data not available.")

    doc.add_page_break()

    # ── evaluation summary ────────────────────────────────────────────────────
    doc.add_heading("Evaluation Summary", level=1)

    doc.add_paragraph(
        f"Overall Score: {eval_scores.get('total_score', 0)}/100  "
        f"(Grade: {eval_scores.get('grade', 'N/A')})"
    ).runs[0].font.bold = True

    doc.add_paragraph(eval_scores.get("overall_recommendation", ""))
    doc.add_paragraph("")

    criteria = ["innovation", "feasibility", "budget_clarity", "impact"]
    for criterion in criteria:
        data = eval_scores.get(criterion, {})
        if not isinstance(data, dict):
            continue

        doc.add_heading(
            f"{criterion.replace('_', ' ').title()} — {data.get('score', 0)}/25",
            level=2
        )
        doc.add_paragraph(f"Strength: {data.get('strength', '')}")
        doc.add_paragraph(f"Weakness: {data.get('weakness', '')}")

    # ── save ─────────────────────────────────────────────────────────────────
    # sanitize topic for filename
    safe_topic = "".join(c if c.isalnum() or c in " _-" else "" for c in topic)
    safe_topic = safe_topic[:40].strip().replace(" ", "_")
    timestamp  = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename   = f"proposal_{safe_topic}_{timestamp}.docx"
    filepath   = os.path.join(EXPORT_DIR, filename)

    doc.save(filepath)
    logger.info(f"Proposal exported to: {filepath}")
    return filepath


def export_proposal_txt(
    proposal_draft: dict,
    budget_timeline: dict,
    eval_scores: dict,
    topic: str
) -> str:
    """
    Plain text fallback export — no dependencies needed.
    Used when python-docx is not installed.

    Args:
        Same as export_proposal_docx

    Returns:
        Path to the saved .txt file
    """
    ensure_export_dir()

    lines = [
        "=" * 70,
        "RESEARCH GRANT PROPOSAL",
        f"Topic: {topic}",
        f"Generated: {datetime.now().strftime('%B %d, %Y')}",
        f"Score: {eval_scores.get('total_score', 0)}/100",
        "=" * 70,
        ""
    ]

    # proposal sections
    for section_name, content in proposal_draft.items():
        lines.append(f"\n{'─' * 50}")
        lines.append(section_name.upper())
        lines.append('─' * 50)
        lines.append(content)

    # budget summary
    lines.append(f"\n{'─' * 50}")
    lines.append("BUDGET SUMMARY")
    lines.append('─' * 50)
    lines.append(f"Total: ${budget_timeline.get('total_cost_usd', 0):,} USD")
    for item in budget_timeline.get("budget_items", []):
        lines.append(
            f"  [{item.get('category')}] {item.get('item')} — "
            f"${item.get('cost_usd', 0):,}"
        )

    # evaluation
    lines.append(f"\n{'─' * 50}")
    lines.append("EVALUATION SCORES")
    lines.append('─' * 50)
    lines.append(f"Total: {eval_scores.get('total_score', 0)}/100 | Grade: {eval_scores.get('grade', 'N/A')}")
    lines.append(eval_scores.get("overall_recommendation", ""))
    for c in ["innovation", "feasibility", "budget_clarity", "impact"]:
        data = eval_scores.get(c, {})
        if isinstance(data, dict):
            lines.append(
                f"  {c.title()}: {data.get('score', 0)}/25  "
                f"[+{data.get('strength', '')}]  "
                f"[-{data.get('weakness', '')}]"
            )

    content_str = "\n".join(lines)

    safe_topic = "".join(c if c.isalnum() or c in " _-" else "" for c in topic)
    safe_topic = safe_topic[:40].strip().replace(" ", "_")
    timestamp  = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename   = f"proposal_{safe_topic}_{timestamp}.txt"
    filepath   = os.path.join(EXPORT_DIR, filename)

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content_str)

    logger.info(f"Proposal exported (txt) to: {filepath}")
    return filepath


def export_scores_json(eval_scores: dict, run_id: int) -> str:
    """
    Export evaluation scores as a JSON file.
    Useful for loading into analysis tools or for your project report.

    Args:
        eval_scores: Full evaluation dict from EvaluationAgent
        run_id: Run ID from SQLite (used in filename)

    Returns:
        Path to the saved .json file
    """
    ensure_export_dir()

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename  = f"evaluation_run{run_id}_{timestamp}.json"
    filepath  = os.path.join(EXPORT_DIR, filename)

    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(eval_scores, f, indent=2)

    logger.info(f"Scores exported to: {filepath}")
    return filepath
